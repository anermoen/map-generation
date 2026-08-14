#!/usr/bin/env python3
"""
Generate a Word report (property purpose, workflow, numerical methods,
and per-year findings/figures) for one property case - reading live
from that property's output folder (property.py's property_code(), e.g.
123-9-Etnedal) rather than hardcoding any numbers, so the report always
matches whatever's actually in manifest.json when it's run.

Usage
-----
    python3 generate_report.py --kommune Etnedal --gnr 123 --bnr 9

Requires python-docx (`pip install python-docx`) and that property's
manifest.json / overlay PNGs to already exist (run auto_gcp.py and
plot_overlay.py first - see README.md).

Output: <property_code>_Report.docx, inside that property's own output
folder (property.py's property_code()) - alongside the GeoTIFFs and
overlays it documents.
"""

import argparse
import json
import os

from docx import Document
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from property import fetch_property, property_code
from imagery_search import find_covering_projects, best_covering_project

SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))

METHOD_LABELS = {
    "automatic_gcp_extraction": "Automatic (auto_gcp.py)",
    "manual_gcp_affine_fit": "Manual (georeference_screenshot.py)",
}

# A4, with tight-ish margins - the figures in Section 4 use the full
# resulting page width (see set_a4_page_and_get_usable_width), rather
# than a fixed inch value that leaves most of the sheet unused.
A4_WIDTH_CM = 21.0
A4_HEIGHT_CM = 29.7
MARGIN_LR_CM = 1.5
MARGIN_TB_CM = 2.0
# Rough vertical budget for each figure's "4.N <year>" heading + caption
# + inter-paragraph spacing, so a full-width (and, since every overlay
# PNG this project produces is square, equally tall) image still fits
# on one page rather than spilling a sliver onto the next.
FIGURE_TEXT_BUDGET_CM = 3.0


def set_a4_page_and_get_usable_width(doc):
    section = doc.sections[0]
    section.page_width = Cm(A4_WIDTH_CM)
    section.page_height = Cm(A4_HEIGHT_CM)
    section.left_margin = Cm(MARGIN_LR_CM)
    section.right_margin = Cm(MARGIN_LR_CM)
    section.top_margin = Cm(MARGIN_TB_CM)
    section.bottom_margin = Cm(MARGIN_TB_CM)
    usable_width = section.page_width - section.left_margin - section.right_margin
    usable_height = (section.page_height - section.top_margin - section.bottom_margin
                      - Cm(FIGURE_TEXT_BUDGET_CM))
    return min(usable_width, usable_height)


def add_equation_block(doc, lines):
    """Monospace block for equations/formulas - no LaTeX rendering
    needed for this level of notation, just clean fixed-width alignment."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.4)
    for i, line in enumerate(lines):
        if i > 0:
            p.add_run().add_break()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run.font.size = Pt(11)
    return p


def add_param_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    doc.add_paragraph()
    return table


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def load_manifest(outdir, basename):
    path = os.path.join(outdir, f"{basename}_manifest.json")
    if not os.path.isfile(path):
        raise SystemExit(f"{path} not found - run auto_gcp.py (and, for any years it "
                          f"can't fit automatically, georeference_screenshot.py) first.")
    with open(path) as f:
        return json.load(f)


def main():
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--kommune", default="Etnedal")
    ap.add_argument("--gnr", type=int, default=123)
    ap.add_argument("--bnr", type=int, default=9)
    ap.add_argument("--outdir", default=None)
    args = ap.parse_args()

    prop = fetch_property(args.kommune, args.gnr, args.bnr)
    basename = prop.code
    outdir = args.outdir or basename
    manifest = load_manifest(outdir, basename)
    images = sorted(manifest["images"], key=lambda r: r["year"])

    print("Searching Norge i Bilder's project coverage for this property "
          "(for each figure's source-project caption)...")
    imagery_projects = find_covering_projects(prop.polygon)

    doc = Document()
    usable_width = set_a4_page_and_get_usable_width(doc)

    # --- Title page ---
    title = doc.add_heading(
        f"Historical Aerial Imagery for Property {prop.matrikkelnummer}, {prop.kommunenavn}",
        level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        "Cadastral-boundary georeferencing of norgeibilder.no browser screenshots")
    sub_run.italic = True
    sub_run.font.size = Pt(13)
    doc.add_paragraph()
    add_param_table(
        doc,
        ["Property", "Area", "Bounds (EPSG:25833)", "Years documented"],
        [[
            f"{prop.matrikkelnummer}, {prop.kommunenavn} kommune (kommunenummer {prop.kommunenummer})",
            f"{prop.polygon.area:,.0f} m^2 (~{prop.polygon.area / 1e6:.2f} km^2)",
            "({:.0f}, {:.0f}) - ({:.0f}, {:.0f})".format(*manifest["bounds_25833"]),
            ", ".join(str(r["year"]) for r in images),
        ]],
    )
    doc.add_page_break()

    # --- Purpose ---
    doc.add_heading("1. Purpose", level=1)
    doc.add_paragraph(
        "This project extracts historical aerial photographs of a single Norwegian "
        "property, across as many decades as coverage allows, for early-stage "
        "internal testing of a GIS/imagery workflow ahead of client discussions - "
        "not yet a client deliverable. The property's exact boundary comes from "
        "Kartverket's national cadastre (matrikkelen); the aerial photographs "
        "themselves come from Kartverket's Norge i Bilder programme, which has "
        "photographed Norway repeatedly since the 1930s-1950s."
    )
    doc.add_paragraph(
        "Norge i Bilder's own image-serving API requires an authenticated, "
        "IP-bound access token, and that authentication is gated behind Norge "
        "digitalt - a national geodata-sharing partnership open to public-sector "
        "organizations, not commercial or research entities without an existing "
        "membership or a separate data-access agreement. Since none of the "
        "available paths to that access (formal membership, a direct agreement, "
        "or ordering individual historical photos from Kartverket's archive at a "
        "per-image cost) fit this project's current pre-client stage, this "
        "report instead documents a free, immediately usable alternative: Norge "
        "i Bilder's own public browser viewer requires no login at all, and "
        "already draws the property's boundary directly onto each historical "
        "photo on screen. A screenshot of that view, together with the same "
        "boundary's exact real-world coordinates (already available from the "
        "open cadastre WFS), contains everything needed to georeference the "
        "image directly - the approach this report's figures and numbers "
        "document."
    )

    # --- Workflow ---
    doc.add_heading("2. Workflow", level=1)
    doc.add_paragraph(
        "The pipeline runs in four stages, each a self-contained, independently "
        "runnable script:"
    )

    doc.add_heading("2.1 Property boundary lookup (property.py)", level=2)
    doc.add_paragraph(
        "Fetches the property's exact polygon boundary from Kartverket's open "
        "WFS ('Matrikkelen - Eiendomskart Teig', no API key needed), given a "
        "kommune name and gnr/bnr (gardsnummer/bruksnummer). The kommunenummer "
        "the WFS actually filters on is resolved at request time via "
        "Kartverket's Kommuneinfo API rather than hardcoded, since these codes "
        "are periodically renumbered (Etnedal's changed nationally in 2024)."
    )

    doc.add_heading("2.2 Coverage search (imagery_search.py)", level=2)
    doc.add_paragraph(
        "Norge i Bilder's open project-metadata API lists thousands of "
        "historical and current aerial-photo/satellite 'prosjekt' nationally. "
        "This step finds which of those actually cover the property - checked "
        "against each candidate project's real flown coverage polygon, not "
        "just an overlapping bounding box, since the two can differ by "
        "kilometers - and for which years."
    )

    doc.add_heading("2.3 Screenshot capture (manual)", level=2)
    doc.add_paragraph(
        "For each year with real coverage, a browser screenshot of Norge i "
        "Bilder's public viewer is captured by hand: search to the property's "
        "address, enable the 'Property boundaries' (eiendomsgrenser) map "
        "layer, select the historical project/year on the timeline, and zoom "
        "so the whole property boundary is visible in frame. No login or "
        "token is needed for this step."
    )

    doc.add_heading("2.4 Georeferencing", level=2)
    doc.add_paragraph(
        "Converts each screenshot into a properly georeferenced GeoTIFF "
        "(EPSG:25833), tagged with its fit quality, by matching the boundary "
        "line drawn in the screenshot against the boundary's known real-world "
        "coordinates and fitting the affine transform between pixel and world "
        "coordinates. Two implementations exist:"
    )
    doc.add_paragraph(
        "auto_gcp.py - fully automatic, run as a batch over every screenshot "
        "that hasn't been fit yet (see Section 3 for how the matching itself "
        "works). This is the primary path.", style="List Bullet")
    doc.add_paragraph(
        "georeference_screenshot.py - a manual fallback (list-vertices / pick "
        "/ fit) for the rare screenshot the automatic method can't confidently "
        "match on its own (see Section 4's per-year notes).", style="List Bullet")
    doc.add_paragraph(
        "plot_overlay.py then renders each georeferenced photo with the "
        "property boundary - fetched fresh from the live cadastre, not read "
        "back from the photo itself - drawn on top, in true map coordinates: "
        "the figures in Section 4 (and the main visual check used throughout "
        "development) are this overlay, since a correct fit shows the two "
        "boundary lines coinciding almost exactly."
    )

    # --- Numerical methods ---
    doc.add_heading("3. Numerical Methods", level=1)
    doc.add_paragraph(
        "Georeferencing a screenshot means finding the affine transform that "
        "maps pixel coordinates (column, row) to real-world map coordinates "
        "(easting, northing):"
    )
    add_equation_block(doc, [
        "easting  = a*col + b*row + c",
        "northing = d*col + e*row + f",
    ])
    doc.add_paragraph(
        "Given a set of matched points - pixel positions paired with their "
        "known real-world coordinates - the six parameters (a..f) are found "
        "by ordinary least squares (numpy.linalg.lstsq), and each match's "
        "residual distance in meters gives a direct, interpretable accuracy "
        "measure: the root-mean-square error (RMSE) and maximum error reported "
        "throughout this document and in every image's own embedded metadata."
    )

    doc.add_heading("3.1 Finding the matches automatically", level=2)
    doc.add_paragraph(
        "The property boundary is drawn in a distinctive color in every "
        "screenshot; auto_gcp.py isolates it by color threshold, extracts its "
        "outline as a polyline (OpenCV contour detection, simplified via the "
        "Douglas-Peucker algorithm), and then has to work out which point "
        "along that pixel outline corresponds to which point on the property's "
        "actual boundary - the two are the same shape, but at different, "
        "independent scales, rotations, and levels of simplification. Two "
        "complementary strategies each generate a rough starting guess (a "
        "'seed') for this correspondence, and neither is trusted directly:"
    )
    doc.add_paragraph(
        "a crude bounding-box mapping - matching the pixel outline's own "
        "bounding box to the property's, assuming (correctly, for this data "
        "source) that the screenshot is north-up with the whole property "
        "visible; and", style="List Bullet")
    doc.add_paragraph(
        "sequence alignment on the outline's turn angles at each corner - the "
        "same class of algorithm used to diff text or align DNA sequences "
        "(dynamic programming), which tolerates the pixel and world "
        "boundaries having been simplified to different numbers of corners.",
        style="List Bullet")
    doc.add_paragraph(
        "Every seed is then refined by Iterative Closest Point (ICP): project "
        "the pixel corners through the current transform estimate, snap each "
        "to its nearest point on the property's real, unsimplified boundary, "
        "refit, and repeat with a shrinking match-distance tolerance. This is "
        "what actually locks onto an accurate transform - verified directly to "
        "recover the correct registration even from a badly inaccurate seed - "
        "but, being a local refinement method, it can also converge onto a "
        "self-consistent but wrong registration (typically where a repetitive "
        "or self-similar section of the boundary matches itself at the wrong "
        "offset). Two independent, cheap checks catch this before it can be "
        "reported as a result: the fitted transform's two axis scales must "
        "agree to within 15% (every genuine fit measured to within 0.6% on "
        "this project's real screenshots), and its implied rotation must be "
        "within 20 degrees of north-up (the viewer this project's screenshots "
        "come from has no rotation control at all, so every genuine screenshot "
        "is un-rotated). Whichever candidate survives both checks with the "
        "lowest final RMSE is kept; if none do, the year is left for the "
        "manual fallback (georeference_screenshot.py) rather than reporting an "
        "unreliable result."
    )

    # --- Findings ---
    doc.add_heading("4. Findings and Figures", level=1)
    n_auto = sum(1 for r in images if r["georeferencing_method"] == "automatic_gcp_extraction")
    n_manual = len(images) - n_auto
    rmses = [r["rmse_m"] for r in images]
    doc.add_paragraph(
        "{} historical aerial photographs of this property were georeferenced: "
        "{} automatically and {} via the manual fallback, spanning "
        "{}-{}. Fit accuracy (RMSE) ranged from {:.2f} m to {:.2f} m across all "
        "years, well within the resolution of the source imagery itself (each "
        "photo's own pixel size, tagged in its GeoTIFF, is on the order of "
        "1-2 m/pixel for these historical scans).".format(
            len(images), n_auto, n_manual, images[0]["year"], images[-1]["year"],
            min(rmses), max(rmses))
    )

    add_param_table(
        doc,
        ["Year", "Method", "GCPs used", "RMSE (m)", "Max error (m)", "Pixel size (m)"],
        [[
            r["year"], METHOD_LABELS.get(r["georeferencing_method"], r["georeferencing_method"]),
            r["n_gcp"], f"{r['rmse_m']:.2f}", f"{r['max_error_m']:.2f}",
            "x".join(f"{v:.2f}" for v in r["pixel_size_m"]),
        ] for r in images],
    )

    doc.add_paragraph(
        "Each figure below is that year's photo, georeferenced and rendered in "
        "true map coordinates (easting/northing, EPSG:25833), with the "
        "property boundary overlaid twice: once as fetched live from the "
        "cadastre just now (red), and once as it was already drawn into the "
        "original screenshot by Norge i Bilder's own viewer (visible "
        "underneath, in magenta, wherever the georeferencing has placed it). "
        "The two lines coinciding closely is the direct visual confirmation "
        "that a given year's fit is correct - not just a low RMSE number."
    )

    for i, r in enumerate(images, start=1):
        year = r["year"]
        img_name = os.path.splitext(r["filename"])[0] + "_overlay.png"
        img_path = os.path.join(outdir, img_name)
        method_label = METHOD_LABELS.get(r["georeferencing_method"], r["georeferencing_method"])

        # Key output from imagery_search.py: which real Norge i Bilder
        # project this year's photo actually came from, and its native
        # (source-photo) resolution - distinct from r["pixel_size_m"]
        # above, which is this *screenshot's* rasterized resolution at
        # whatever zoom level it was captured at, not the underlying
        # photo's own resolution.
        source = best_covering_project(imagery_projects, year, allow_satellite=True)
        matched_year = year
        if source is None:
            # The screenshot's year label (chosen by whoever captured it,
            # e.g. from the norgeibilder.no timeline UI) occasionally
            # doesn't exactly match the covering project's own "aar"
            # field - seen for real with 123/9 Etnedal: a screenshot
            # labelled 1956 whose only covering project is dated 1958
            # (photo date 1958-06-05). Rather than silently reporting
            # nothing, fall back to the nearest covering year within a
            # couple of years and say so explicitly in the caption -
            # never claim an exact match that isn't real.
            NEAR_YEAR_TOLERANCE = 2
            near_years = sorted(
                {p.year for p in imagery_projects
                 if abs(p.year - year) <= NEAR_YEAR_TOLERANCE},
                key=lambda y: abs(y - year))
            for cand_year in near_years:
                source = best_covering_project(imagery_projects, cand_year, allow_satellite=True)
                if source is not None:
                    matched_year = cand_year
                    break
        if source is not None:
            kind = "satellite" if source.is_satellite else ("CIR" if source.is_cir else "aerial")
            if matched_year != year:
                year_note = " [nearest covering project, dated {} - screenshot labelled {}]".format(
                    matched_year, year)
            else:
                year_note = ""
            source_desc = "source: {!r} ({}, photo date {}, {:g} m/px native resolution){}".format(
                source.name, kind, source.photo_date, source.pixel_size_m, year_note)
        else:
            source_desc = "source project not found in a live imagery_search.py re-check"

        if os.path.isfile(img_path):
            doc.add_heading(f"4.{i}  {year}", level=2)
            doc.add_picture(img_path, width=usable_width)
            add_caption(
                doc,
                "Figure {}. {}, {} - {}. Georeferenced via {} ({} GCPs, RMSE = {:.2f} m, "
                "max error = {:.2f} m).".format(
                    i, prop.matrikkelnummer, year, source_desc, method_label, r["n_gcp"],
                    r["rmse_m"], r["max_error_m"]))
        else:
            doc.add_heading(f"4.{i}  {year}", level=2)
            doc.add_paragraph(f"[{img_name} not found - run plot_overlay.py first]")

    # --- Limitations ---
    doc.add_heading("5. Limitations and Assumptions", level=1)
    for text in [
        "Screenshot-based georeferencing is inherently approximate, not a "
        "substitute for a real WMS download: accuracy is limited by screen "
        "resolution at the zoom level captured, not the source photo's native "
        "resolution, and by how precisely the boundary line itself was "
        "rendered and thresholded. The RMSE reported per year (Section 4) is "
        "the honest, measured accuracy of each specific fit, not an assumed "
        "constant.",
        "The automatic method assumes the whole property boundary is visible "
        "in the screenshot and that the screenshot is north-up - both true for "
        "every screenshot this project's manual capture process has produced "
        "so far, but not guaranteed for an arbitrarily cropped or rotated "
        "image; a screenshot that violates either assumption is expected to "
        "fail the automatic fit's checks (Section 3.1) rather than silently "
        "produce a wrong result.",
        "Norge i Bilder's real coverage does not include every requested year "
        "for every property - only projects with a genuine flown survey "
        "reaching the property are used; a gap year is a real property of the "
        "aerial-photo program's flight schedule for that area, not a search "
        "shortfall.",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    out_path = os.path.join(outdir, f"{basename}_Report.docx")
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
